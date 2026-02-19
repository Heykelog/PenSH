from fastapi import FastAPI, Depends, HTTPException, status, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from sqlalchemy.orm import Session
from sqlalchemy import create_engine
from typing import List, Optional
import os
import uuid
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from io import BytesIO
from urllib.parse import quote

# Optional DOCX support
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except Exception:
    Document = None
    WD_ALIGN_PARAGRAPH = None
    DOCX_AVAILABLE = False

# Register fonts early
try:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    pdfmetrics.registerFont(TTFont('DejaVuSerif', '/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf'))
    pdfmetrics.registerFont(TTFont('DejaVuSerif-Bold', '/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf'))
    pdfmetrics.registerFontFamily('DejaVuSerif', normal='DejaVuSerif', bold='DejaVuSerif-Bold', italic=None, boldItalic=None)
    print("DejaVu fonts registered successfully in main.py")
except Exception as e:
    print(f"Font registration failed in main.py: {e}")

from database import get_db, engine, Base
from models import Report, Finding, OwaspTemplate, KnowledgeBaseTemplate, RiskLevel, OwaspCategory, POCImage, Customer, Tester
import schemas
from pdf_generator import PentestReportGenerator
from owasp_data import OWASP_TOP_10_TEMPLATES
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment

# Create tables
Base.metadata.create_all(bind=engine)

app = FastAPI(
    title="Pentest Şablon Yönetim Sistemi",
    description="OWASP Top 10 tabanlı penetrasyon testi rapor yönetim uygulaması",
    version="1.0.0"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, specify exact origins
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create directories
os.makedirs("uploads", exist_ok=True)
os.makedirs("static", exist_ok=True)
os.makedirs("reports", exist_ok=True)

# Static files
app.mount("/static", StaticFiles(directory="static"), name="static")
app.mount("/uploads", StaticFiles(directory="uploads"), name="uploads")

# Initialize OWASP templates on startup
@app.on_event("startup")
async def initialize_owasp_templates():
    db = next(get_db())
    
    # Check if templates already exist
    existing_count = db.query(OwaspTemplate).count()
    if existing_count == 0:
        # Add OWASP Top 10 templates
        for category, template_data in OWASP_TOP_10_TEMPLATES.items():
            owasp_template = OwaspTemplate(
                category=category,
                title=template_data["title"],
                description=template_data["description"],
                impact=template_data["impact"],
                solution=template_data["solution"],
                risk_level=template_data["risk_level"]
            )
            db.add(owasp_template)
        
        db.commit()
        print("OWASP Top 10 şablonları başarıyla yüklendi.")
    
    db.close()

# Health check
@app.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

# Reports endpoints
@app.get("/reports", response_model=List[schemas.Report])
async def get_reports(skip: int = 0, limit: int = 100, db: Session = Depends(get_db)):
    """Tüm raporları listele"""
    reports = db.query(Report).offset(skip).limit(limit).all()
    return reports

@app.get("/reports/{report_id}", response_model=schemas.Report)
async def get_report(report_id: int, db: Session = Depends(get_db)):
    """Belirli bir raporu getir"""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Rapor bulunamadı")
    return report

@app.post("/reports", response_model=schemas.Report, status_code=status.HTTP_201_CREATED)
async def create_report(report: schemas.ReportCreate, db: Session = Depends(get_db)):
    """Yeni rapor oluştur"""
    db_report = Report(**report.dict())
    db.add(db_report)
    db.commit()
    db.refresh(db_report)
    return db_report

@app.put("/reports/{report_id}", response_model=schemas.Report)
async def update_report(report_id: int, report_update: schemas.ReportUpdate, db: Session = Depends(get_db)):
    """Rapor güncelle"""
    db_report = db.query(Report).filter(Report.id == report_id).first()
    if not db_report:
        raise HTTPException(status_code=404, detail="Rapor bulunamadı")
    
    update_data = report_update.dict(exclude_unset=True)
    for field, value in update_data.items():
        setattr(db_report, field, value)
    
    db.commit()
    db.refresh(db_report)
    return db_report

@app.post("/reports/{report_id}/upload-logo")
async def upload_report_logo(
    report_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """Rapor için logo yükle"""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Rapor bulunamadı")
    
    # Validate file type
    if not file.content_type.startswith('image/'):
        raise HTTPException(status_code=400, detail="Sadece resim dosyaları yüklenebilir")
    
    # Create logos directory
    logos_dir = os.path.join("backend", "static", "logos")
    os.makedirs(logos_dir, exist_ok=True)
    
    # Generate unique filename
    file_extension = os.path.splitext(file.filename)[1]
    unique_filename = f"logo_report_{report_id}_{uuid.uuid4()}{file_extension}"
    
    # Save file
    file_path = os.path.join(logos_dir, unique_filename)
    with open(file_path, "wb") as buffer:
        content = await file.read()
        buffer.write(content)
    
    # Update report
    report.logo_path = file_path
    db.commit()
    db.refresh(report)
    
    return {
        "message": "Logo başarıyla yüklendi",
        "logo_path": file_path,
        "filename": unique_filename
    }

@app.delete("/reports/{report_id}")
async def delete_report(report_id: int, db: Session = Depends(get_db)):
    """Rapor sil.

    Not: Eğer bu rapora ait bulgulardan üretilmiş Bilgi Bankası şablonları varsa,
    şablonlar silinmez; sadece ilgili bulgularla olan ilişkileri koparılır.
    Böylece rapor ve bulgular silinirken şablonlar Bilgi Bankası'nda kalmaya devam eder.
    """
    db_report = db.query(Report).filter(Report.id == report_id).first()
    if not db_report:
        raise HTTPException(status_code=404, detail="Rapor bulunamadı")

    # Bu rapora ait bulguların ID'lerini al
    finding_ids = [f.id for f in db_report.findings]

    if finding_ids:
        # İlgili Bilgi Bankası şablonlarının finding_id alanını boşalt
        db.query(KnowledgeBaseTemplate).filter(
            KnowledgeBaseTemplate.finding_id.in_(finding_ids)
        ).update({"finding_id": None}, synchronize_session=False)

    db.delete(db_report)
    db.commit()
    return {"message": "Rapor başarıyla silindi"}

# Findings endpoints
@app.get("/findings", response_model=List[schemas.Finding])
async def get_findings(
    skip: int = 0, 
    limit: int = 100, 
    report_id: Optional[int] = None,
    risk_level: Optional[RiskLevel] = None,
    owasp_category: Optional[OwaspCategory] = None,
    db: Session = Depends(get_db)
):
    """Bulguları listele (filtreleme seçenekleri ile)"""
    query = db.query(Finding)
    
    if report_id:
        query = query.filter(Finding.report_id == report_id)
    if risk_level:
        query = query.filter(Finding.risk_level == risk_level)
    if owasp_category:
        query = query.filter(Finding.owasp_category == owasp_category)
    
    findings = query.offset(skip).limit(limit).all()
    return findings

@app.get("/findings/{finding_id}", response_model=schemas.Finding)
async def get_finding(finding_id: int, db: Session = Depends(get_db)):
    """Belirli bir bulguyu getir"""
    finding = db.query(Finding).filter(Finding.id == finding_id).first()
    if not finding:
        raise HTTPException(status_code=404, detail="Bulgu bulunamadı")
    return finding

@app.post("/findings", response_model=schemas.Finding, status_code=status.HTTP_201_CREATED)
async def create_finding(finding: schemas.FindingCreate, db: Session = Depends(get_db)):
    """Yeni bulgu oluştur"""
    # Check if report exists
    report = db.query(Report).filter(Report.id == finding.report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Rapor bulunamadı")
    
    db_finding = Finding(**finding.dict())
    db.add(db_finding)
    db.commit()
    db.refresh(db_finding)
    return db_finding

@app.put("/findings/{finding_id}", response_model=schemas.Finding)
async def update_finding(finding_id: int, finding_update: schemas.FindingUpdate, db: Session = Depends(get_db)):
    """Bulgu güncelle"""
    db_finding = db.query(Finding).filter(Finding.id == finding_id).first()
    if not db_finding:
        raise HTTPException(status_code=404, detail="Bulgu bulunamadı")
    
    update_data = finding_update.dict(exclude_unset=True)
    for field, value in update_data.items():
        setattr(db_finding, field, value)
    
    db.commit()
    db.refresh(db_finding)
    return db_finding

@app.delete("/findings/{finding_id}")
async def delete_finding(finding_id: int, db: Session = Depends(get_db)):
    """Bulgu sil.

    Not: Eğer bu bulgudan oluşturulmuş Bilgi Bankası şablonları varsa,
    şablonlar silinmez; sadece bu bulgu ile olan ilişkisi koparılır.
    Böylece bulgu rapordan silinirken şablonlar Bilgi Bankası'nda kalmaya devam eder.
    """
    db_finding = db.query(Finding).filter(Finding.id == finding_id).first()
    if not db_finding:
        raise HTTPException(status_code=404, detail="Bulgu bulunamadı")

    # Bu bulgudan üretilmiş Bilgi Bankası şablonlarının ilişki alanını boşalt
    db.query(KnowledgeBaseTemplate).filter(
        KnowledgeBaseTemplate.finding_id == finding_id
    ).update({"finding_id": None})

    db.delete(db_finding)
    db.commit()
    return {"message": "Bulgu başarıyla silindi"}

@app.post("/findings/reorder")
async def reorder_findings(
    reorder_data: dict,  # { "report_id": int, "orderedIds": [id1, id2, ...] }
    db: Session = Depends(get_db)
):
    """Bulguların sırasını güncelle"""
    report_id = reorder_data.get("report_id")
    ordered_ids = reorder_data.get("orderedIds", [])
    
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Rapor bulunamadı")
    
    # Her bulgu için display_order güncelle
    for index, finding_id in enumerate(ordered_ids):
        finding = db.query(Finding).filter(Finding.id == finding_id).first()
        if finding and finding.report_id == report_id:
            finding.display_order = index
    
    db.commit()
    return {"message": "Bulgular başarıyla sıralandı"}

# OWASP Templates endpoints
@app.get("/owasp-templates", response_model=List[schemas.OwaspTemplate])
async def get_owasp_templates(db: Session = Depends(get_db)):
    """OWASP Top 10 şablonlarını listele"""
    templates = db.query(OwaspTemplate).filter(OwaspTemplate.is_active == True).all()
    return templates

@app.get("/owasp-templates/{template_id}", response_model=schemas.OwaspTemplate)
async def get_owasp_template(template_id: int, db: Session = Depends(get_db)):
    """Belirli bir OWASP şablonunu getir"""
    template = db.query(OwaspTemplate).filter(OwaspTemplate.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="OWASP şablonu bulunamadı")
    return template

@app.post("/findings/from-owasp-template", response_model=schemas.Finding, status_code=status.HTTP_201_CREATED)
async def create_finding_from_owasp_template(
    template_id: int,
    report_id: int,
    custom_title: Optional[str] = None,
    custom_affected_area: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """OWASP şablonundan bulgu oluştur"""
    # Check if report exists
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Rapor bulunamadı")
    
    # Get OWASP template
    template = db.query(OwaspTemplate).filter(OwaspTemplate.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="OWASP şablonu bulunamadı")
    
    # Create finding from template
    db_finding = Finding(
        title=custom_title or template.title,
        description=template.description,
        affected_area=custom_affected_area,
        risk_level=template.risk_level,
        owasp_category=template.category,
        solution=template.solution,
        impact=template.impact,
        report_id=report_id
    )
    
    db.add(db_finding)
    db.commit()
    db.refresh(db_finding)
    return db_finding

@app.post("/findings/from-knowledge-base-template", response_model=schemas.Finding, status_code=status.HTTP_201_CREATED)
async def create_finding_from_knowledge_base_template(
    template_id: int,
    report_id: int,
    custom_title: Optional[str] = None,
    custom_affected_area: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """Bilgi bankasından (Knowledge Base template'inden) bulgu oluştur"""
    # Check if report exists
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Rapor bulunamadı")
    
    # Get Knowledge Base template
    template = db.query(KnowledgeBaseTemplate).filter(KnowledgeBaseTemplate.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Bilgi bankası şablonu bulunamadı")
    
    # Create finding from template (transfer all template data)
    db_finding = Finding(
        title=custom_title or template.title,
        description=template.description,
        affected_area=custom_affected_area or template.affected_area,
        risk_level=template.risk_level,
        owasp_category=template.owasp_category,
        solution=template.solution,
        steps_to_reproduce=template.steps_to_reproduce,
        impact=template.impact,
        request=template.request,
        response=template.response,
        cvss_score=template.cvss_score,
        cwe_id=template.cwe_id,
        refs=template.refs,
        report_id=report_id
    )
    
    db.add(db_finding)
    db.commit()
    db.refresh(db_finding)
    return db_finding

# Knowledge Base Templates endpoints
@app.get("/knowledge-base-templates", response_model=List[schemas.KnowledgeBaseTemplate])
async def get_knowledge_base_templates(db: Session = Depends(get_db)):
    """Bilgi Bankası şablonlarını listele (Kullanıcı tarafından eklenenler)"""
    templates = db.query(KnowledgeBaseTemplate).filter(
        KnowledgeBaseTemplate.is_from_finding == True
    ).all()
    return templates

@app.post("/knowledge-base-templates", response_model=schemas.KnowledgeBaseTemplate, status_code=status.HTTP_201_CREATED)
async def create_knowledge_base_template(
    template: schemas.KnowledgeBaseTemplateCreate,
    db: Session = Depends(get_db)
):
    """Bulgudan Bilgi Bankası şablonu oluştur"""
    if template.finding_id:
        finding = db.query(Finding).filter(Finding.id == template.finding_id).first()
        if not finding:
            raise HTTPException(status_code=404, detail="Bulgu bulunamadı")
    
    db_template = KnowledgeBaseTemplate(
        title=template.title,
        description=template.description,
        affected_area=template.affected_area,
        impact=template.impact,
        solution=template.solution,
        risk_level=template.risk_level,
        owasp_category=template.owasp_category,
        steps_to_reproduce=template.steps_to_reproduce,
        cvss_score=template.cvss_score,
        cwe_id=template.cwe_id,
        refs=template.refs,
        is_from_finding=True,
        finding_id=template.finding_id
    )
    
    db.add(db_template)
    db.commit()
    db.refresh(db_template)
    return db_template

@app.get("/knowledge-base-templates/{template_id}", response_model=schemas.KnowledgeBaseTemplate)
async def get_knowledge_base_template(template_id: int, db: Session = Depends(get_db)):
    """Belirli bir Bilgi Bankası şablonunu getir"""
    template = db.query(KnowledgeBaseTemplate).filter(KnowledgeBaseTemplate.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Bilgi Bankası şablonu bulunamadı")
    return template

@app.put("/knowledge-base-templates/{template_id}", response_model=schemas.KnowledgeBaseTemplate)
async def update_knowledge_base_template(
    template_id: int,
    template_data: schemas.KnowledgeBaseTemplateCreate,
    db: Session = Depends(get_db)
):
    """Bilgi Bankası şablonunu güncelle"""
    template = db.query(KnowledgeBaseTemplate).filter(KnowledgeBaseTemplate.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Bilgi Bankası şablonu bulunamadı")
    
    # Update fields
    template.title = template_data.title
    template.description = template_data.description
    template.affected_area = template_data.affected_area
    template.impact = template_data.impact
    template.solution = template_data.solution
    template.risk_level = template_data.risk_level
    template.owasp_category = template_data.owasp_category
    template.steps_to_reproduce = template_data.steps_to_reproduce
    template.request = template_data.request
    template.response = template_data.response
    template.cvss_score = template_data.cvss_score
    template.cwe_id = template_data.cwe_id
    template.refs = template_data.refs
    
    db.commit()
    db.refresh(template)
    return template

@app.delete("/knowledge-base-templates/{template_id}")
async def delete_knowledge_base_template(template_id: int, db: Session = Depends(get_db)):
    """Bilgi Bankası şablonunu sil"""
    template = db.query(KnowledgeBaseTemplate).filter(KnowledgeBaseTemplate.id == template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail="Bilgi Bankası şablonu bulunamadı")
    
    db.delete(template)
    db.commit()
    return {"message": "Bilgi Bankası şablonu başarıyla silindi"}

@app.post("/findings/{finding_id}/save-to-knowledge-base", response_model=schemas.KnowledgeBaseTemplate, status_code=status.HTTP_201_CREATED)
async def save_finding_to_knowledge_base(
    finding_id: int,
    db: Session = Depends(get_db)
):
    """Bulguyu Bilgi Bankası şablonuna ekle"""
    finding = db.query(Finding).filter(Finding.id == finding_id).first()
    if not finding:
        raise HTTPException(status_code=404, detail="Bulgu bulunamadı")
    
    # Bilgi Bankası şablonunu oluştur
    kb_template = KnowledgeBaseTemplate(
        title=finding.title,
        description=finding.description,
        affected_area=finding.affected_area,
        impact=finding.impact,
        solution=finding.solution,
        risk_level=finding.risk_level,
        owasp_category=finding.owasp_category,
        steps_to_reproduce=finding.steps_to_reproduce,
        request=getattr(finding, 'request', None),
        response=getattr(finding, 'response', None),
        cvss_score=finding.cvss_score,
        cwe_id=finding.cwe_id,
        refs=finding.refs,
        is_from_finding=True,
        finding_id=finding_id
    )
    
    db.add(kb_template)
    db.commit()
    db.refresh(kb_template)
    return kb_template

# PDF Export endpoint
@app.post("/export/pdf/{report_id}")
async def export_report_to_pdf(
    report_id: int,
    include_logo: bool = False,
    db: Session = Depends(get_db)
):
    """Raporu PDF formatında dışa aktar"""
    # Get report with findings
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Rapor bulunamadı")
    
    # Generate unique filename
    filename = f"pentest_report_{report_id}_{uuid.uuid4().hex[:8]}.pdf"
    output_path = os.path.join("reports", filename)
    
    # Get findings for the report, sorted by display_order
    findings = db.query(Finding).filter(Finding.report_id == report_id).order_by(
        Finding.display_order.asc(),
        Finding.id.asc()
    ).all()
    
    # Generate PDF
    try:
        print(f"PDF oluşturuluyor: Report ID {report_id}, Findings: {len(findings)}")
        generator = PentestReportGenerator()
        output_path = generator.generate_report(report, findings, output_path)
        
        print(f"PDF başarıyla oluşturuldu: {output_path}")
        
        # Create clean filename from report title
        clean_title = "".join(c for c in report.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        clean_title = clean_title.replace(' ', '_')
        clean_filename = f"{clean_title}_pentest_report.pdf"
        
        return FileResponse(
            output_path,
            media_type="application/pdf",
            filename=clean_filename
        )
    except Exception as e:
        print(f"PDF oluşturma hatası: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"PDF oluşturma hatası: {str(e)}")

# Excel Export endpoint
@app.post("/export/xlsx/{report_id}")
async def export_report_to_xlsx(
    report_id: int,
    db: Session = Depends(get_db)
):
    """Raporu Excel (XLSX) formatında dışa aktar"""
    # Get report with findings
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Rapor bulunamadı")

    findings = db.query(Finding).filter(Finding.report_id == report_id).all()

    try:
        wb = Workbook()

        # Styling helpers
        header_fill = PatternFill(start_color="FFEDF2FF", end_color="FFEDF2FF", fill_type="solid")
        header_font = Font(bold=True)
        center = Alignment(horizontal="center")
        wrap = Alignment(wrap_text=True, vertical="top")

        # Sheet 1: Report Summary
        ws_summary = wb.active
        ws_summary.title = "Rapor Özeti"
        ws_summary.append(["Alan", "Değer"])
        ws_summary["A1"].font = header_font
        ws_summary["B1"].font = header_font
        ws_summary["A1"].fill = header_fill
        ws_summary["B1"].fill = header_fill

        summary_rows = [
            ("Rapor Başlığı", report.title),
            ("Müşteri", report.client_name or (report.customer.name if getattr(report, "customer", None) else "")),
            ("Test Tarihi", report.test_date or ""),
            ("Test Uzmanı", report.tester_name or (report.tester.name if getattr(report, "tester", None) else "")),
            ("Oluşturulma", report.created_at.strftime("%d/%m/%Y %H:%M")),
        ]
        for row in summary_rows:
            ws_summary.append(list(row))
        ws_summary.column_dimensions['A'].width = 25
        ws_summary.column_dimensions['B'].width = 80

        # Sheet 2: Scope & Methodology
        ws_meta = wb.create_sheet(title="Kapsam & Metodoloji")
        ws_meta.append(["Kapsam"])
        ws_meta["A1"].font = header_font
        ws_meta["A1"].fill = header_fill
        ws_meta.append([report.scope or ""])
        ws_meta.append([])
        ws_meta.append(["Metodoloji"])
        ws_meta["A4"].font = header_font
        ws_meta["A4"].fill = header_fill
        ws_meta.append([report.methodology or ""])
        ws_meta.column_dimensions['A'].width = 120
        for row in ws_meta.iter_rows(min_row=2, max_row=ws_meta.max_row, min_col=1, max_col=1):
            for cell in row:
                cell.alignment = wrap

        # Sheet 3: Findings
        ws_findings = wb.create_sheet(title="Bulgular")
        headers = [
            "ID", "Başlık", "Risk", "OWASP", "Etkilenen Alan", "Açıklama", "Çözüm",
            "Tekrarlama Adımları", "Etki", "CVSS", "CWE", "Referanslar", "Oluşturulma"
        ]
        ws_findings.append(headers)
        for col_idx in range(1, len(headers) + 1):
            cell = ws_findings.cell(row=1, column=col_idx)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center

        for f in findings:
            ws_findings.append([
                f.id,
                f.title,
                f.risk_level.value if hasattr(f.risk_level, 'value') else f.risk_level,
                f.owasp_category.value if f.owasp_category else "",
                f.affected_area or "",
                f.description,
                f.solution or "",
                f.steps_to_reproduce or "",
                f.impact or "",
                f.cvss_score or "",
                f.cwe_id or "",
                f.refs or "",
                f.created_at.strftime("%d/%m/%Y %H:%M") if f.created_at else "",
            ])

        # Column widths and wrapping
        widths = [8, 40, 12, 28, 30, 80, 80, 80, 40, 20, 16, 40, 20]
        for idx, width in enumerate(widths, start=1):
            ws_findings.column_dimensions[chr(64 + idx) if idx <= 26 else f"A{chr(64 + idx - 26)}"].width = width
        for row in ws_findings.iter_rows(min_row=2, max_row=ws_findings.max_row):
            for cell in row:
                cell.alignment = wrap

        # Save to file
        filename = f"pentest_report_{report_id}_{uuid.uuid4().hex[:8]}.xlsx"
        output_path = os.path.join("reports", filename)
        wb.save(output_path)

        # Clean filename for download
        clean_title = "".join(c for c in report.title if c.isalnum() or c in (' ', '-', '_')).rstrip().replace(' ', '_')
        download_name = f"{clean_title}_pentest_report.xlsx"

        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            filename=download_name
        )
    except Exception as e:
        print(f"XLSX oluşturma hatası: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"XLSX oluşturma hatası: {str(e)}")

# File upload endpoint for logo
@app.post("/upload/logo")
async def upload_logo(file: UploadFile = File(...)):
    """Logo dosyası yükle"""
    if not file.content_type.startswith("image/"):
        raise HTTPException(status_code=400, detail="Sadece resim dosyaları kabul edilir")
    
    # Save logo
    logo_path = "static/logo.png"
    with open(logo_path, "wb") as buffer:
        content = await file.read()
        buffer.write(content)
    
    return {"message": "Logo başarıyla yüklendi", "path": logo_path}

# Search endpoints
@app.get("/search/findings")
async def search_findings(
    q: str,
    report_id: Optional[int] = None,
    db: Session = Depends(get_db)
):
    """Bulgularda arama yap"""
    query = db.query(Finding)
    
    if report_id:
        query = query.filter(Finding.report_id == report_id)
    
    # Search in title, description, and solution
    search_filter = (
        Finding.title.contains(q) |
        Finding.description.contains(q) |
        Finding.solution.contains(q)
    )
    
    findings = query.filter(search_filter).all()
    return findings

@app.get("/search/reports")
async def search_reports(q: str, db: Session = Depends(get_db)):
    """Raporlarda arama yap"""
    search_filter = (
        Report.title.contains(q) |
        Report.description.contains(q) |
        Report.client_name.contains(q)
    )
    
    reports = db.query(Report).filter(search_filter).all()
    return reports

# Statistics endpoint
@app.get("/statistics")
async def get_statistics(db: Session = Depends(get_db)):
    """İstatistik bilgilerini getir"""
    total_reports = db.query(Report).count()
    total_findings = db.query(Finding).count()
    
    # Risk level distribution
    risk_distribution = {}
    for risk_level in RiskLevel:
        count = db.query(Finding).filter(Finding.risk_level == risk_level).count()
        risk_distribution[risk_level.value] = count
    
    # OWASP category distribution
    owasp_distribution = {}
    for category in OwaspCategory:
        count = db.query(Finding).filter(Finding.owasp_category == category).count()
        owasp_distribution[category.value] = count
    
    return {
        "total_reports": total_reports,
        "total_findings": total_findings,
        "risk_distribution": risk_distribution,
        "owasp_distribution": owasp_distribution
    }

# POC Image endpoints
@app.post("/findings/{finding_id}/poc-images")
async def upload_poc_image(
    finding_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """Bulguya POC ekran görüntüsü yükle"""
    # Check if finding exists
    finding = db.query(Finding).filter(Finding.id == finding_id).first()
    if not finding:
        raise HTTPException(status_code=404, detail="Finding not found")
    
    # Validate file type
    if not file.content_type.startswith('image/'):
        raise HTTPException(status_code=400, detail="Only image files are allowed")
    
    # Generate unique filename
    file_extension = os.path.splitext(file.filename)[1]
    unique_filename = f"{uuid.uuid4()}{file_extension}"
    
    # Create upload directory for this finding
    finding_upload_dir = os.path.join("uploads", "poc_images", str(finding_id))
    os.makedirs(finding_upload_dir, exist_ok=True)
    
    # Save file
    file_path = os.path.join(finding_upload_dir, unique_filename)
    with open(file_path, "wb") as buffer:
        content = await file.read()
        buffer.write(content)
    
    # Save to database
    poc_image = POCImage(
        filename=unique_filename,
        original_filename=file.filename,
        file_path=file_path,
        file_size=len(content),
        mime_type=file.content_type,
        finding_id=finding_id
    )
    
    db.add(poc_image)
    db.commit()
    db.refresh(poc_image)
    
    return poc_image

@app.get("/findings/{finding_id}/poc-images")
async def get_poc_images(finding_id: int, db: Session = Depends(get_db)):
    """Bulguya ait POC ekran görüntülerini getir"""
    finding = db.query(Finding).filter(Finding.id == finding_id).first()
    if not finding:
        raise HTTPException(status_code=404, detail="Finding not found")
    
    poc_images = db.query(POCImage).filter(POCImage.finding_id == finding_id).all()
    return poc_images

@app.delete("/poc-images/{image_id}")
async def delete_poc_image(image_id: int, db: Session = Depends(get_db)):
    """POC ekran görüntüsünü sil"""
    poc_image = db.query(POCImage).filter(POCImage.id == image_id).first()
    if not poc_image:
        raise HTTPException(status_code=404, detail="POC image not found")
    
    # Delete file from filesystem
    if os.path.exists(poc_image.file_path):
        os.remove(poc_image.file_path)
    
    # Delete from database
    db.delete(poc_image)
    db.commit()
    
    return {"message": "POC image deleted successfully"}

@app.get("/poc-images/{image_id}/download")
async def download_poc_image(image_id: int, db: Session = Depends(get_db)):
    """POC ekran görüntüsünü indir"""
    poc_image = db.query(POCImage).filter(POCImage.id == image_id).first()
    if not poc_image:
        raise HTTPException(status_code=404, detail="POC image not found")
    
    if not os.path.exists(poc_image.file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        poc_image.file_path,
        media_type=poc_image.mime_type,
        filename=poc_image.original_filename
    )

# Excel Export endpoint
@app.get("/export/xlsx/{report_id}")
async def export_report_to_xlsx(
    report_id: int,
    db: Session = Depends(get_db)
):
    """Raporu Excel formatında dışa aktar (çok sayfalı)."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Rapor bulunamadı")

    findings = db.query(Finding).filter(Finding.report_id == report_id).all()

    wb = Workbook()
    ws_summary = wb.active
    ws_summary.title = "Özet"

    # Headers
    ws_summary.append(["Rapor Başlığı", report.title])
    ws_summary.append(["Müşteri", report.client_name or "-"])
    ws_summary.append(["Test Tarihi", report.test_date or "-"])
    ws_summary.append(["Test Uzmanı", report.tester_name or "-"])
    ws_summary.append(["Oluşturulma", report.created_at.strftime("%Y-%m-%d %H:%M") if isinstance(report.created_at, datetime) else str(report.created_at)])
    ws_summary.append([])
    ws_summary.append(["Toplam Bulgu", len(findings)])

    # Findings sheet
    ws_findings = wb.create_sheet("Bulgular")
    header = [
        "ID", "Başlık", "Risk", "OWASP", "Etkilenen Alan", "Açıklama", "Etki", "Çözüm", "CWE", "CVSS", "Oluşturulma"
    ]
    ws_findings.append(header)
    bold = Font(bold=True)
    for cell in ws_findings[1]:
        cell.font = bold
        cell.alignment = Alignment(vertical="top")

    for f in findings:
        ws_findings.append([
            f.id,
            f.title,
            f.risk_level.value if hasattr(f.risk_level, 'value') else f.risk_level,
            f.owasp_category.value if getattr(f, 'owasp_category', None) and hasattr(f.owasp_category, 'value') else (f.owasp_category or ""),
            f.affected_area or "",
            f.description or "",
            f.impact or "",
            f.solution or "",
            f.cwe_id or "",
            f.cvss_score or "",
            f.created_at.strftime("%Y-%m-%d %H:%M") if isinstance(f.created_at, datetime) else str(f.created_at),
        ])

    # Auto width (simple)
    for sheet in [ws_summary, ws_findings]:
        for column_cells in sheet.columns:
            length = max(len(str(cell.value)) if cell.value is not None else 0 for cell in column_cells)
            sheet.column_dimensions[column_cells[0].column_letter].width = min(max(length + 2, 12), 60)

    # Optional extra sheets similar to sample usage
    ws_scope = wb.create_sheet("Kapsam")
    ws_scope.append(["Test Kapsamı"])
    ws_scope["A1"].font = bold
    ws_scope.append([report.scope or "-"])

    ws_method = wb.create_sheet("Metodoloji")
    ws_method.append(["Test Metodolojisi"])
    ws_method["A1"].font = bold
    ws_method.append([report.methodology or "-"])

    # VULNERABILITIES sheet with requested fields
    ws_vuln = wb.create_sheet("VULNERABILITIES")
    vuln_headers = [
        "Name (Required)",
        "IP (Required)",
        "Hostname",
        "Virtual Host",
        "Description (Required)",
        "Solution",
        "Owasp",
        "Date(yyyy-MM-dd HH:mm)",
        "Vulnerability Description",
    ]
    ws_vuln.append(vuln_headers)
    for cell in ws_vuln[1]:
        cell.font = bold
        cell.alignment = Alignment(vertical="top")

    # Map findings to the VULNERABILITIES sheet columns
    for f in findings:
        # Derive fields
        name_val = f.title or ""
        ip_val = ""  # Not tracked; left blank intentionally
        hostname_val = ""  # Not tracked; left blank intentionally
        virtual_host_val = (f.affected_area or "")
        description_val = f.description or ""
        solution_val = f.solution or ""
        owasp_val = (
            f.owasp_category.value
            if getattr(f, 'owasp_category', None) and hasattr(f.owasp_category, 'value')
            else (f.owasp_category or "")
        )
        date_val = (
            f.created_at.strftime("%Y-%m-%d %H:%M")
            if isinstance(f.created_at, datetime)
            else str(f.created_at)
        )
        vuln_desc_val = f.impact or f.description or ""

        ws_vuln.append([
            name_val,
            ip_val,
            hostname_val,
            virtual_host_val,
            description_val,
            solution_val,
            owasp_val,
            date_val,
            vuln_desc_val,
        ])

    # Stream to memory
    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)

    clean_title = "".join(c for c in report.title if c.isalnum() or c in (' ', '-', '_')).rstrip().replace(' ', '_')
    filename = f"{clean_title}_pentest_report.xlsx"

    ascii_fallback = filename.encode('ascii', 'ignore').decode('ascii') or "report.xlsx"
    content_disp = f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{quote(filename)}"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": content_disp}
    )

# DOCX Export endpoint
@app.get("/export/docx/{report_id}")
async def export_report_to_docx(
    report_id: int,
    db: Session = Depends(get_db)
):
    """Raporu Word (DOCX) formatında dışa aktar."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Rapor bulunamadı")

    # Get findings ordered by display_order (same as in PDF)
    findings = db.query(Finding).filter(Finding.report_id == report_id).order_by(
        Finding.display_order.asc(), Finding.id.asc()
    ).all()

    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=500, detail="DOCX desteği yüklü değil")

    doc = Document()

    # Title page
    title = doc.add_heading(report.title or "Pentest Raporu", 0)
    if WD_ALIGN_PARAGRAPH:
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_p = doc.add_paragraph()
    info_runs = [
        ("Müşteri: ", True, report.client_name or (report.customer.name if getattr(report, "customer", None) else "-")),
        ("\nTest Tarihi: ", True, report.test_date or "-"),
        ("\nTest Uzmanı: ", True, report.tester_name or (report.tester.name if getattr(report, "tester", None) else "-")),
        ("\nRapor Tarihi: ", True, datetime.now().strftime("%d/%m/%Y")),
        ("\nToplam Bulgu: ", True, str(len(findings)))
    ]
    for text, bold, value in info_runs:
        r = info_p.add_run(text)
        r.bold = bold
        info_p.add_run(value)

    # ---- Executive Summary ----
    doc.add_page_break()
    doc.add_heading("YÖNETİCİ ÖZETİ", level=1)

    risk_counts = {}
    for f in findings:
        key = f.risk_level.value if hasattr(f.risk_level, 'value') else str(f.risk_level)
        risk_counts[key.lower()] = risk_counts.get(key.lower(), 0) + 1

    summary = doc.add_paragraph()
    summary.add_run("Rapor Özeti:\n").bold = True
    summary.add_run(
        f"Bu rapor, {report.client_name or 'Müşteri'} için gerçekleştirilen penetrasyon testi sonuçlarını içerir. "
        f"Toplam {len(findings)} bulgu tespit edilmiştir.\n"
    )
    summary.add_run(f"Test Tarihi: ").bold = True; summary.add_run(f"{report.test_date or '-'}\n")
    summary.add_run(f"Test Uzmanı: ").bold = True; summary.add_run(f"{report.tester_name or '-'}\n")

    # ---- Findings Summary Table ----
    doc.add_heading("BULGULAR ÖZETİ", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "#"
    hdr_cells[1].text = "Bulgu Başlığı"
    hdr_cells[2].text = "Risk Seviyesi"
    hdr_cells[3].text = "OWASP Kategorisi"

    for idx, f in enumerate(findings, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = f.title or ""
        row[2].text = (f.risk_level.value if hasattr(f.risk_level, 'value') else str(f.risk_level)).upper()
        row[3].text = (f.owasp_category.value if getattr(f, 'owasp_category', None) and hasattr(f.owasp_category, 'value') else (f.owasp_category or ""))

    # ---- Detailed Findings ----
    doc.add_page_break()
    doc.add_heading("DETAYLI BULGULAR", level=1)

    for i, f in enumerate(findings, start=1):
        doc.add_heading(f"{i}. {f.title}", level=2)

        # Risk level and metadata
        meta = doc.add_paragraph()
        meta.add_run("Risk: ").bold = True
        meta.add_run((f.risk_level.value if hasattr(f.risk_level, 'value') else str(f.risk_level)) or "-")
        if getattr(f, 'owasp_category', None):
            meta.add_run("  |  OWASP: ").bold = True
            meta.add_run(f.owasp_category.value if hasattr(f.owasp_category, 'value') else str(f.owasp_category))
        if f.affected_area:
            meta.add_run("  |  Etkilenen Alan: ").bold = True
            meta.add_run(f.affected_area)
        if f.cvss_score:
            meta.add_run("  |  CVSS: ").bold = True
            meta.add_run(f.cvss_score)
        if f.cwe_id:
            meta.add_run("  |  CWE: ").bold = True
            meta.add_run(f.cwe_id)

        # Description
        if f.description:
            doc.add_heading("Açıklama", level=3)
            doc.add_paragraph(f.description)
        
        # Impact
        if f.impact:
            doc.add_heading("Etki", level=3)
            doc.add_paragraph(f.impact)
        
        # Steps to reproduce
        if f.steps_to_reproduce:
            doc.add_heading("Adım Adım (Step-by-Step)", level=3)
            doc.add_paragraph(f.steps_to_reproduce)
        
        # Request/Response Examples
        if f.request or f.response:
            doc.add_heading("Request/Response Örneği", level=3)
            
            if f.request:
                req_p = doc.add_paragraph()
                req_p.add_run("Request:\n").bold = True
                req_code = doc.add_paragraph(f.request, style='No Spacing')
                req_code.style = 'No Spacing'
            
            if f.response:
                resp_p = doc.add_paragraph()
                resp_p.add_run("Response:\n").bold = True
                resp_code = doc.add_paragraph(f.response, style='No Spacing')
                resp_code.style = 'No Spacing'
        
        # Solution
        if f.solution:
            doc.add_heading("Çözüm", level=3)
            doc.add_paragraph(f.solution)
        
        # References
        if f.refs:
            doc.add_heading("Referanslar", level=3)
            doc.add_paragraph(f.refs)

        if i < len(findings):
            doc.add_page_break()

    # ---- Methodology ----
    doc.add_heading("TEST METODOLOJİSİ", level=1)
    if report.methodology:
        doc.add_paragraph(report.methodology)

    # ---- Recommendations ----
    doc.add_heading("ÖNERİLER VE SONUÇLAR", level=1)
    doc.add_paragraph(
        "Kritik ve yüksek riskli açıkların öncelikli olarak giderilmesi, düzenli test ve yama yönetimi, "
        "izleme ve loglama süreçlerinin güçlendirilmesi önerilir."
    )

    # Stream document
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)

    clean_title = "".join(c for c in (report.title or "rapor") if c.isalnum() or c in (' ', '-', '_')).rstrip().replace(' ', '_')
    filename = f"{clean_title}_pentest_report.docx"

    ascii_fallback = filename.encode('ascii', 'ignore').decode('ascii') or "report.docx"
    content_disp = f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{quote(filename)}"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": content_disp}
    )

# Customer Management endpoints
@app.get("/customers", response_model=List[schemas.Customer])
async def get_customers(db: Session = Depends(get_db)):
    """Müşterileri listele"""
    customers = db.query(Customer).order_by(Customer.name).all()
    return customers

@app.post("/customers", response_model=schemas.Customer, status_code=status.HTTP_201_CREATED)
async def create_customer(customer: schemas.CustomerCreate, db: Session = Depends(get_db)):
    """Yeni müşteri oluştur"""
    # If this is set as default, unset all other defaults
    if customer.is_default:
        db.query(Customer).update({"is_default": False})
    
    db_customer = Customer(**customer.dict())
    db.add(db_customer)
    db.commit()
    db.refresh(db_customer)
    return db_customer

@app.put("/customers/{customer_id}", response_model=schemas.Customer)
async def update_customer(customer_id: int, customer_update: schemas.CustomerUpdate, db: Session = Depends(get_db)):
    """Müşteri güncelle"""
    db_customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not db_customer:
        raise HTTPException(status_code=404, detail="Müşteri bulunamadı")
    
    # If this is set as default, unset all other defaults
    if customer_update.is_default:
        db.query(Customer).update({"is_default": False})
    
    update_data = customer_update.dict(exclude_unset=True)
    for field, value in update_data.items():
        setattr(db_customer, field, value)
    
    db.commit()
    db.refresh(db_customer)
    return db_customer

@app.delete("/customers/{customer_id}")
async def delete_customer(customer_id: int, db: Session = Depends(get_db)):
    """Müşteri sil"""
    db_customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not db_customer:
        raise HTTPException(status_code=404, detail="Müşteri bulunamadı")
    
    db.delete(db_customer)
    db.commit()
    return {"message": "Müşteri başarıyla silindi"}

@app.get("/customers/default", response_model=schemas.Customer)
async def get_default_customer(db: Session = Depends(get_db)):
    """Varsayılan müşteriyi getir"""
    customer = db.query(Customer).filter(Customer.is_default == True).first()
    if not customer:
        # Return first customer if no default is set
        customer = db.query(Customer).first()
    return customer

# Tester Management endpoints
@app.get("/testers", response_model=List[schemas.Tester])
async def get_testers(db: Session = Depends(get_db)):
    """Test uzmanlarını listele"""
    testers = db.query(Tester).order_by(Tester.name).all()
    return testers

@app.post("/testers", response_model=schemas.Tester, status_code=status.HTTP_201_CREATED)
async def create_tester(tester: schemas.TesterCreate, db: Session = Depends(get_db)):
    """Yeni test uzmanı oluştur"""
    # If this is set as default, unset all other defaults
    if tester.is_default:
        db.query(Tester).update({"is_default": False})
    
    db_tester = Tester(**tester.dict())
    db.add(db_tester)
    db.commit()
    db.refresh(db_tester)
    return db_tester

@app.put("/testers/{tester_id}", response_model=schemas.Tester)
async def update_tester(tester_id: int, tester_update: schemas.TesterUpdate, db: Session = Depends(get_db)):
    """Test uzmanı güncelle"""
    db_tester = db.query(Tester).filter(Tester.id == tester_id).first()
    if not db_tester:
        raise HTTPException(status_code=404, detail="Test uzmanı bulunamadı")
    
    # If this is set as default, unset all other defaults
    if tester_update.is_default:
        db.query(Tester).update({"is_default": False})
    
    update_data = tester_update.dict(exclude_unset=True)
    for field, value in update_data.items():
        setattr(db_tester, field, value)
    
    db.commit()
    db.refresh(db_tester)
    return db_tester

@app.delete("/testers/{tester_id}")
async def delete_tester(tester_id: int, db: Session = Depends(get_db)):
    """Test uzmanı sil"""
    db_tester = db.query(Tester).filter(Tester.id == tester_id).first()
    if not db_tester:
        raise HTTPException(status_code=404, detail="Test uzmanı bulunamadı")
    
    db.delete(db_tester)
    db.commit()
    return {"message": "Test uzmanı başarıyla silindi"}

@app.get("/testers/default", response_model=schemas.Tester)
async def get_default_tester(db: Session = Depends(get_db)):
    """Varsayılan test uzmanını getir"""
    tester = db.query(Tester).filter(Tester.is_default == True).first()
    if not tester:
        # Return first tester if no default is set
        tester = db.query(Tester).first()
    return tester

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
